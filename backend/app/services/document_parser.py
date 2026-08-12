"""文档解析服务：将上传文件提取为纯文本。

支持格式：PDF / Word(.docx) / Excel(.xlsx) / 纯文本 / Markdown / 常见源代码文件。
代码文件原样保留源码文本（含注释/文档字符串），交由分块器按顶层定义切分，
支持「代码文件入库 → 代码问答」。
"""
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

# 基础文档格式
BASE_EXTS = {"pdf", "docx", "xlsx", "txt", "md"}
# 可解析的源代码扩展名（原样读取源码）
CODE_EXTS = {
    "py", "js", "ts", "tsx", "jsx", "java", "go", "cpp", "c",
    "h", "hpp", "cs", "rs", "rb", "php", "swift", "kt",
}
SUPPORTED_EXTS = BASE_EXTS | CODE_EXTS

# 扩展名 -> chunker 语言 key（用于按顶层定义切块；无映射的扩展名走普通分块）
CODE_LANG = {
    "py": "python",
    "js": "javascript",
    "ts": "typescript",
    "tsx": "typescript",
    "jsx": "javascript",
    "java": "java",
    "go": "go",
    "cpp": "cpp",
    "h": "cpp",
    "hpp": "cpp",
    "c": "c",
    "cs": "csharp",
    "rs": "rust",
    "rb": "ruby",
    "php": "php",
    "swift": "swift",
    "kt": "kotlin",
}


def parse_document(filename: str, path: Path) -> str:
    """按扩展名分发到对应解析器，返回提取后的纯文本。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"不支持的文件类型: .{ext}")
    if ext == "pdf":
        return _parse_pdf(path)
    if ext == "docx":
        return _parse_docx(path)
    if ext == "xlsx":
        return _parse_xlsx(path)
    if ext in CODE_EXTS:
        return _parse_code(path)
    # txt / md 直接读取（尝试多种编码，保证中文不乱码）
    return _parse_text(path)


def _parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_xlsx(path: Path) -> str:
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## 工作表: {ws.title}")
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
        parts.append("\n".join(rows))
    wb.close()
    return "\n\n".join(parts)


def _parse_text(path: Path) -> str:
    # 依次尝试 UTF-8 / GBK / latin-1，避免中文编码问题
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="latin-1", errors="ignore")


def _parse_code(path: Path) -> str:
    """源代码文件：原样读取源码（注释 / 文档字符串一并保留）。

    代码文件大概率 UTF-8（现代编辑器默认），保留编码容错路径防旧项目 GBK。
    """
    return _parse_text(path)
